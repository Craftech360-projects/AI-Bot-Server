# create_sample_docx.py

from docx import Document
import os

# Directory to save the sample files
sample_dir = "sample_data"
os.makedirs(sample_dir, exist_ok=True)

# Sample content
doc = Document()
doc.add_heading('Sample Document', 0)

doc.add_paragraph(
    "This is a sample DOCX document. It contains information that can be used to test the chatbot's ability to retrieve and process content."
)
doc.add_paragraph(
    "For example, this document might describe the basics of natural language processing (NLP) or provide an overview of machine learning algorithms."
)

# Save the DOCX file
doc.save(os.path.join(sample_dir, "sample_document.docx"))

print(f"DOCX sample created in the '{sample_dir}' directory.")
